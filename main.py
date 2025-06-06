from pdf2docx import Converter
from docx import Document
from docx2pdf import convert
import requests
import sys

# 1. Convert PDF to DOCX
def convert_pdf_to_docx(pdf_path, docx_path):
    print("Converting PDF to DOCX...")
    try:
        # First attempt - with strict image handling disabled
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None, pages=None, 
                  ignore_images=True)
        cv.close()
        print("PDF converted (without images)")
    except Exception as e:
        print(f"Standard conversion failed: {str(e)}")
        
        try:
            # Second attempt - try a more drastic approach with PyMuPDF directly
            import fitz  # PyMuPDF
            import io
            from docx import Document
            
            print("Attempting text-only extraction...")
            doc = Document()
            pdf = fitz.open(pdf_path)
            
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                text = page.get_text()
                
                # Skip empty pages
                if text.strip():
                    # Add page number as heading
                    doc.add_heading(f"Page {page_num + 1}", level=1)
                    
                    # Split text by newlines and add as paragraphs
                    for line in text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line)
            
            # Save the document
            doc.save(docx_path)
            pdf.close()
            print("Text-only extraction successful")
            
        except Exception as e2:
            print(f"All conversion attempts failed: {str(e2)}")
            raise RuntimeError("Could not convert PDF to DOCX. The PDF may have unsupported features.")

# 2. Translate text using Ollama
def translate_text_ollama(text, model="mistral"):
    if not text.strip():
        return text
    
    print(f"\rSending translation request...", end='', flush=True)
    prompt = f"Translate the following text from Finnish to English. Just give the best translated sentences without quotes. Do not give alternatives or other comments.\n\n{text}"
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": model, "prompt": prompt, "stream": False},
            timeout=60
        )
        result = response.json().get("response", "").strip()
        return result if result else text
    except requests.exceptions.Timeout:
        print(f"\nTranslation request timed out. Is Ollama running?")
        return text
    except Exception as e:
        print(f"\nTranslation failed: {e}")
        return text

# 3. Translate DOCX content
def translate_docx(input_docx, output_docx, model="gemma3:4b"):
    print("Translating DOCX content...")
    doc = Document(input_docx)
    
    # Get total number of paragraphs for progress calculation
    total_paragraphs = len(doc.paragraphs)
    processed = 0
    
    # Skip empty paragraphs count
    non_empty_paragraphs = sum(1 for para in doc.paragraphs if para.text.strip())
    print(f"Found {total_paragraphs} paragraphs ({non_empty_paragraphs} non-empty)")
    
    for para in doc.paragraphs:
        original = para.text
        
        # Only translate non-empty paragraphs
        if original.strip():
            translated = translate_text_ollama(original, model=model)
            
            # Preserve formatting by keeping the first run's formatting
            # and replacing all runs with a single one
            if para.runs:
                # Store the formatting attributes we want to preserve
                first_run = para.runs[0]
                
                # Clear all runs
                for run in para.runs[:]:
                    para._p.remove(run._r)
                
                # Create a new run with the translated text and original formatting
                new_run = para.add_run(translated)
                
                # Copy formatting from the first run
                new_run.bold = first_run.bold
                new_run.italic = first_run.italic
                new_run.underline = first_run.underline
                new_run.font.size = first_run.font.size
                new_run.font.name = first_run.font.name
                new_run.font.color.rgb = first_run.font.color.rgb
            else:
                # If there are no runs, just set the text
                para.text = translated
        
        # Update progress bar
        processed += 1
        percentage = (processed / total_paragraphs) * 100
        bar_length = 40
        filled_length = int(bar_length * processed // total_paragraphs)
        bar = '█' * filled_length + '-' * (bar_length - filled_length)
        print(f"\rProgress: |{bar}| {percentage:.1f}% ({processed}/{total_paragraphs})", end='', flush=True)
        
        # Save after each paragraph to preserve work
        if processed % 10 == 0:
            doc.save(output_docx)
    
    print("\nTranslation completed!")
    doc.save(output_docx)

# 4. Convert DOCX to PDF
def convert_docx_to_pdf(docx_path, pdf_path):
    print("Converting translated DOCX to PDF...")
    convert(docx_path, pdf_path)


# ==== Main workflow ====
def main():
    # Source path and model name (check carefully)
    pdf_file = "raw/document.pdf"
    model = "gemma3:4b"
    
    # Check if filename was provided as command-line argument
    if len(sys.argv) > 1:
        pdf_file = sys.argv[1]
        print(f"Using provided PDF file: {pdf_file}")
    else:
        print(f"Using default PDF file: {pdf_file}")
        
    # Other paths
    docx_file = "raw/temp.docx"
    translated_docx = "raw/translated.docx"
    translated_pdf = "raw/translated.pdf"

    convert_pdf_to_docx(pdf_file, docx_file)
    translate_docx(docx_file, translated_docx, model)
    convert_docx_to_pdf(translated_docx, translated_pdf)
    print("✅ Translation complete: translated.pdf")

if __name__ == "__main__":
    main()
